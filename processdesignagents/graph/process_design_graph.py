import os
from pathlib import Path
import json
import pypandoc

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None

from typing import Dict, Any

from langchain_openai import ChatOpenAI
# from langchain_anthropic import ChatAnthropic
from langchain_google_genai import ChatGoogleGenerativeAI

from langgraph.prebuilt import ToolNode

from dotenv import load_dotenv

from processdesignagents.default_config import DEFAULT_CONFIG
from processdesignagents.agents.utils.json_tools import (
    convert_risk_json_to_markdown,
)
from processdesignagents.agents.utils.agent_sizing_tools import (
    size_heat_exchanger_basic,
    size_pump_basic
)
from processdesignagents.agents.utils.equipment_stream_markdown import (
    equipments_and_streams_dict_to_markdown,
)

from .setup import GraphSetup
from .propagator import Propagator
from langgraph.checkpoint.memory import MemorySaver

load_dotenv()

class ProcessDesignGraph:
    """Main class that orchestractes the process design workflow."""
    
    def __init__(
        self,
        debug=False,
        config: Dict[str, Any]=None,
        delay_time: float = 0.5,
    ):
        """Initialize the process design agents graph and component.
        Args:
            debug: Whether to run in debug mode
            config: Configuration dictionary
        """
        self.debug = debug
        self.config = config or DEFAULT_CONFIG
        # Initialize LLMs
        # if self.config["llm_provider"].lower() == "openai" or self.config["llm_provider"] == "ollama" or self.config["llm_provider"] == "openrouter":
        if self.config["llm_provider"] == "openrouter":
            base_url = self.get_url_by_name(self.config["llm_provider"].lower())
            api_key = os.getenv("OPENROUTER_API_KEY")
            self.deep_thinking_llm = ChatOpenAI(model=self.config["deep_think_llm"], base_url=base_url, api_key=api_key)
            self.quick_thinking_llm = ChatOpenAI(model=self.config["quick_think_llm"], base_url=base_url, api_key=api_key)
        elif self.config["llm_provider"] == "ollama":
            base_url = self.get_url_by_name(self.config["llm_provider"].lower())
            self.deep_thinking_llm = ChatOpenAI(model=self.config["deep_think_llm"], base_url=base_url)
            self.quick_thinking_llm = ChatOpenAI(model=self.config["quick_think_llm"], base_url=base_url)
        # elif self.config["llm_provider"].lower() == "anthropic":
        #     self.deep_thinking_llm = ChatAnthropic(model=self.config["deep_think_llm"], base_url=self.config["backend_url"])
        #     self.quick_thinking_llm = ChatAnthropic(model=self.config["quick_think_llm"], base_url=self.config["backend_url"])
        elif self.config["llm_provider"].lower() == "google":
            self.deep_thinking_llm = ChatGoogleGenerativeAI(model=self.config["deep_think_llm"])
            self.quick_thinking_llm = ChatGoogleGenerativeAI(model=self.config["quick_think_llm"])
        else:
            raise ValueError(f"Unsupported LLM provider: {self.config['llm_provider']}")

        self.deep_thinking_llm.temperature = self.config["deep_think_temperature"]
        self.quick_thinking_llm.temperature = self.config["quick_think_temperature"]
        
        self.checkpointer = MemorySaver()
        
        # Create tool nodes
        self.tool_nodes = self._create_tool_nodes()

        self.graph_setup = GraphSetup(
            quick_thinking_llm=self.quick_thinking_llm,
            deep_thinking_llm=self.deep_thinking_llm,
            tool_nodes=self.tool_nodes,
            checkpointer=self.checkpointer,
            delay_time=delay_time,
        )
        
        self.propagator = Propagator()
        self.problem_statement = None
        self.log_state_dict = {}
        
        # Set up the graph
        self.graph = self.graph_setup.setup_graph()
        
        self.graph.get_graph().print_ascii()
        
    def _create_tool_nodes(self) -> Dict[str, ToolNode]:
        """Create tool nodes for different equipment using abstract methods."""
        return {
            "equipment_sizing": ToolNode(
                [
                    size_heat_exchanger_basic,
                    size_pump_basic
                ]
            ),
        }
        
    def get_url_by_name(self, name: str) -> str:
        """
        Retrieves the URL associated with a given name from a predefined list of base URLs.

        Args:
            name (str): The name of the API provider (e.g., 'OpenAI', 'Anthropic').

        Returns:
            str | None: The corresponding URL if found, otherwise None.
        """
        BASE_URLS = [
            ("openai", "https://api.openai.com/v1"),
            ("anthropic", "https://api.anthropic.com/"),
            ("google", "https://generativelanguage.googleapis.com/v1"),
            ("openrouter", "https://openrouter.ai/api/v1"),
            ("ollama", "http://localhost:11434/v1"),        
        ]
        
        for provider, url in BASE_URLS:
            if provider.lower() == name.lower():
                return url
        return None
        
    def propagate(
        self,
        problem_statement: str = "",
        save_markdown: str | None = None,
        save_word_doc: str | None = None,
        manual_concept_selection: bool = False,
    ):
        """Run the agent graph for a problem statement.

        Args:
            problem_statement: Design brief to analyse.
            save_markdown: Optional path for saving the aggregated markdown report.
            manual_concept_selection: When True, prompt the user to choose a concept
                instead of automatically selecting the highest feasibility score.
        """
        self.problem_statement = problem_statement
        previous_provider = self.graph_setup.concept_selection_provider
        if manual_concept_selection:
            def _prompt_user(concept_options):
                print("\nSelect a concept for detailed development:", flush=True)
                for index, option in enumerate(concept_options, start=1):
                    score = option["score"]
                    score_text = f"{score}" if score is not None else "N/A"
                    print(f"{index}. {option['title']} (Feasibility Score: {score_text})", flush=True)
                choice = input("Enter choice (press Enter to use highest score): ").strip()
                if not choice:
                    return None
                try:
                    selection = int(choice)
                except ValueError:
                    print("Invalid selection. Defaulting to highest score.", flush=True)
                    return None
                if 1 <= selection <= len(concept_options):
                    return selection - 1
                print("Selection out of range. Defaulting to highest score.", flush=True)
                return None

            self.graph_setup.concept_selection_provider = _prompt_user
        else:
            # self.graph_setup.concept_selection_provider = None
            def _auto_provider(concept_options):
                return None

            self.graph_setup.concept_selection_provider = _auto_provider


        init_agent_state = self.propagator.create_initial_state(problem_statement)

        args = self.propagator.get_graph_args()
        
        try:
            if self.debug:
                # Debug mode with tracing
                trace = []
                for chunk in self.graph.astream(init_agent_state, **args):
                    if len(chunk["messages"]) == 0:
                        pass
                    else:
                        chunk["messages"][-1].pretty_print()
                        trace.append(chunk)
                final_state = trace[-1]
            else:
                # Run the graph
                final_state = self.graph.invoke(init_agent_state, **args)
                print(f"\n=========================== Finish Line ===========================", flush=True)
        finally:
            self.graph_setup.concept_selection_provider = previous_provider
        
        # Store current state for reflection
        self.curr_state = final_state
        
        # Log state
        self._log_state(final_state)

        if save_markdown:
            self._write_markdown_report(final_state, save_markdown)
            
        if save_word_doc:
            self._write_word_report(final_state, save_word_doc)
        
        return final_state
    
    def _log_state(self, final_state):
        """Log the final state to a JSON file."""
        self.log_state_dict = {
            "problem_statement": final_state.get("problem_statement", ""),
            "requirements": final_state.get("requirements", ""),
            "research_concepts": final_state.get("research_concepts", ""),
            "selected_concept_name": final_state.get("selected_concept_name", ""),
            "selected_concept_details": final_state.get("selected_concept_details", ""),
            "design_basis": final_state.get("design_basis", ""),
            "basic_pfd": final_state.get("basic_pfd", ""),
            "stream_list_template": final_state.get("stream_list_template", ""),
            "stream_list_results": final_state.get("stream_list_results", ""),
            "equipment_list_template": final_state.get("equipment_list_template", ""),
            "equipment_list_results": final_state.get("equipment_list_results", ""),
            "equipment_and_stream_list": final_state.get("equipment_and_stream_list", ""),
            "safety_risk_analyst_report": final_state.get("safety_risk_analyst_report", ""),
            "project_manager_report": final_state.get("project_manager_report", ""),
        }
        
        # Save to file
        directory = Path(f"eval_results/ProcessDesignAgents_logs/")
        directory.mkdir(parents=True, exist_ok=True)
        
        with open(
            f"eval_results/ProcessDesignAgents_logs/full_states_log.json", "w"
        ) as f:
            json.dump(self.log_state_dict, f, indent=4)
        
    def _compose_report_sections(self, final_state: Dict[str, Any]) -> list[tuple[str, str]]:
        raw_equipment_and_streams = final_state.get("equipment_and_stream_list", "")
        if isinstance(raw_equipment_and_streams, str):
            raw_equipment_and_streams = json.loads(raw_equipment_and_streams)
            equipment_and_streams_markdown, _, _ = equipments_and_streams_dict_to_markdown(raw_equipment_and_streams)
        else:
            equipment_and_streams_markdown = ""

        safety_markdown = ""
        if final_state.get("safety_risk_analyst_report"):
            safety_markdown = final_state["safety_risk_analyst_report"]
        sections = [
            ("Problem Statement", final_state.get("problem_statement", "")),
            ("Process Requirements", final_state.get("requirements", "")),
            ("Concept Detail", final_state.get("selected_concept_details", "")),
            ("Design Basis", final_state.get("design_basis", "")),
            ("Basic Process Flow Diagram", final_state.get("basic_pfd", "")),
            ("Equipment and Streams List", equipment_and_streams_markdown),
            ("Safety & Risk Assessment", safety_markdown or final_state.get("safety_risk_analyst_report", "")),
            ("Project Manager Report", final_state.get("project_manager_report", "")),
        ]
        return sections

    def _write_markdown_report(self, final_state: Dict[str, Any], filename: str) -> None:
        sections = self._compose_report_sections(final_state)

        output_lines: list[str] = []
        for title, content in sections:
            if not content:
                continue
            output_lines.append(f"# {title}")
            output_lines.append(content.strip())
            output_lines.append("")  # Blank line separator

        report_text = "\n".join(output_lines).rstrip() + "\n"

        output_path = Path(filename)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(report_text, encoding="utf-8")

    def _write_word_report(self, final_state: Dict[str, Any], filename: str) -> None:
        if Document is None:
            raise ImportError(
                "python-docx is required to export Word reports. Install with `pip install python-docx`."
            )

        sections = self._compose_report_sections(final_state)
        document = Document()

        output_text = ""
        for title, content in sections:
            if not content:
                continue
            # document.add_heading(title, level=1)
            # document.add_paragraph(content.strip())
            output_text += f"# {title}\n{content.strip()}\n\n"

        output_path = Path(filename)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        # document.save(output_path)
        self._export_markdown_to_word(output_text, output_path)

    def _export_markdown_to_word(self, markdown_string: str, output_filename: str = "report.docx"):
        """
        Converts a Markdown string into a formatted Word (.docx) document.

        Args:
            markdown_string: The string containing your final markdown text.
            output_filename: The name of the Word file to create (e.g., "my_document.docx").
        """
        try:
            # Use pypandoc to convert the string
            # 'md' is the source format (Markdown)
            # 'docx' is the target format (Word)
            pypandoc.convert_text(
                markdown_string,
                'docx',
                format='md',
                outputfile=output_filename,
                extra_args=[
                    f"--reference-doc={self.config.get('save_dir')}/template.docx",
                ]
            )
            
            print(f"\nSuccessfully exported Word document to: {os.path.abspath(output_filename)}\n")

        except FileNotFoundError:
            print("\n--- ERROR ---")
            print("Pandoc executable not found.")
            print("Please ensure Pandoc is installed on your system and available in your PATH.")
        except Exception as e:
            print(f"\nAn error occurred during conversion: {e}")
